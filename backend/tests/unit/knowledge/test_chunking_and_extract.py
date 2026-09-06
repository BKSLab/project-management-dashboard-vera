from io import BytesIO
from types import SimpleNamespace
from unittest.mock import AsyncMock

import docx
import pytest
import xlwt
from openpyxl import Workbook

from src.clients.vision import DisabledVisionCapability
from src.knowledge.chunking import chunk_markdown, chunk_text
from src.knowledge.extract import extract_indexable_text


def build_workbook_bytes(build) -> bytes:
    """Собирает книгу в память, применив к листу переданную настройку."""
    workbook = Workbook()
    build(workbook.active)
    stream = BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def build_legacy_workbook_bytes(rows: list[list], title: str = "Смета") -> bytes:
    """Собирает книгу формата .xls, которого openpyxl не умеет записывать."""
    workbook = xlwt.Workbook()
    sheet = workbook.add_sheet(title)
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            if value is not None:
                sheet.write(row_index, column_index, value)
    stream = BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def build_image_bytes() -> bytes:
    """Возвращает минимальную PNG-подобную фикстуру без графических библиотек."""
    return b"\x89PNG\r\n\x1a\nvision-fixture"


def build_pdf_bytes(text: str = "Project plan") -> bytes:
    """Собирает минимальный одностраничный PDF только средствами stdlib."""
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        (
            b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\nendobj\n"
        ),
        b"4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
        (
            f"5 0 obj\n<< /Length {len(stream)} >>\nstream\n".encode("ascii")
            + stream
            + b"\nendstream\nendobj\n"
        ),
    ]
    content = b"%PDF-1.4\n"
    offsets: list[int] = []
    for item in objects:
        offsets.append(len(content))
        content += item
    xref_offset = len(content)
    content += f"xref\n0 {len(objects) + 1}\n".encode("ascii")
    content += b"0000000000 65535 f \n"
    content += b"".join(f"{offset:010} 00000 n \n".encode("ascii") for offset in offsets)
    content += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    ).encode("ascii")
    return content


def test_chunking_keeps_headings_and_splits_deterministically() -> None:
    """Разбиение markdown сохраняет заголовки секций и делит длинный текст одинаково при каждом запуске."""

    content = "# Контекст\nПервый раздел.\n\n## Решение\nВторой раздел."

    chunks = chunk_markdown(content, target_chars=100, overlap_chars=10)

    assert [chunk.heading for chunk in chunks] == ["Контекст", "Решение"]
    assert [chunk.index for chunk in chunks] == [0, 1]
    assert chunks[1].text == "Второй раздел."

    text = " ".join(f"слово-{index}" for index in range(60))

    first = chunk_text(text, target_chars=100, overlap_chars=20)
    second = chunk_text(text, target_chars=100, overlap_chars=20)

    assert first == second
    assert len(first) > 1
    assert all(chunk.strip() for chunk in first)


async def test_extracts_text_from_office_documents() -> None:
    """Извлечение текста из docx, pdf и csv в устаревшей кодировке; неподдерживаемое вложение пропускается."""

    document = docx.Document()
    document.add_paragraph("Описание проекта")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Риск"
    table.cell(0, 1).text = "Срок"
    stream = BytesIO()
    document.save(stream)

    extracted = await extract_indexable_text(
        "паспорт.docx",
        stream.getvalue(),
        vision=DisabledVisionCapability(),
    )

    assert extracted is not None
    assert "Описание проекта" in extracted
    assert "Риск | Срок" in extracted

    extracted = await extract_indexable_text(
        "plan.pdf",
        build_pdf_bytes(),
        vision=DisabledVisionCapability(),
    )

    assert extracted == "Project plan"

    content = "этап;срок\nпроектирование;март".encode("cp1251")

    extracted = await extract_indexable_text("план.csv", content, vision=DisabledVisionCapability())

    assert extracted == "этап;срок\nпроектирование;март"

    assert await extract_indexable_text(
        "архив.zip",
        b"binary",
        vision=DisabledVisionCapability(),
    ) is None


async def test_extracts_spreadsheets_with_their_quirks() -> None:
    """Таблицы: подписанные значения, скрытые строки, объединённые ячейки, смещённая шапка, пустой лист, устаревший xls и его пропуски, битый файл."""

    def build(sheet):
        sheet.title = "Смета"
        sheet.append(["Работа", "Стоимость"])
        sheet.append(["Монтаж", 1500])

    extracted = await extract_indexable_text(
        "смета.xlsx",
        build_workbook_bytes(build),
        vision=DisabledVisionCapability(),
    )

    assert extracted == "Лист: Смета\nРабота: Монтаж, Стоимость: 1500"

    def build(sheet):
        sheet.title = "Этапы"
        sheet.append(["Этап", "Задача"])
        sheet.append(["Подготовка", "Смета"])
        sheet.append([None, "График"])
        sheet.append(["Черновик", "Не считать"])
        sheet.merge_cells("A2:A3")
        sheet.row_dimensions[4].hidden = True

    extracted = await extract_indexable_text(
        "этапы.xlsx",
        build_workbook_bytes(build),
        vision=DisabledVisionCapability(),
    )

    assert extracted == (
        "Лист: Этапы\nЭтап: Подготовка, Задача: Смета\nЭтап: Подготовка, Задача: График"
    )

    def build(sheet):
        sheet.append(["Работа", "Стоимость"])

    extracted = await extract_indexable_text(
        "пусто.xlsx",
        build_workbook_bytes(build),
        vision=DisabledVisionCapability(),
    )

    assert extracted == ""

    def build(sheet):
        sheet.title = "Реестр"
        sheet["A4"] = "Позиция"
        sheet["A5"] = "Кабель"
        sheet["A6"] = "Черновик"
        sheet.row_dimensions[6].hidden = True

    extracted = await extract_indexable_text(
        "реестр.xlsx",
        build_workbook_bytes(build),
        vision=DisabledVisionCapability(),
    )

    assert extracted == "Лист: Реестр\nПозиция: Кабель"

    content = build_legacy_workbook_bytes([["Работа", "Стоимость"], ["Монтаж", 1500]])

    extracted = await extract_indexable_text("смета.xls", content, vision=DisabledVisionCapability())

    assert extracted == "Лист: Смета\nРабота: Монтаж, Стоимость: 1500"

    content = build_legacy_workbook_bytes(
        [["Этап", "Задача"], ["Подготовка", "Смета"], [None, "График"]],
        title="Этапы",
    )

    extracted = await extract_indexable_text("этапы.xls", content, vision=DisabledVisionCapability())

    assert extracted == (
        "Лист: Этапы\nЭтап: Подготовка, Задача: Смета\nЭтап: Подготовка, Задача: График"
    )

    with pytest.raises(ValueError):
        await extract_indexable_text(
            "битая.xlsx",
            b"not a workbook at all",
            vision=DisabledVisionCapability(),
        )


async def test_image_is_recognised_only_with_vision_model() -> None:
    """С включённой моделью зрения изображение распознаётся, без неё пропускается."""
    # Изображение уходит в способность распознавания как есть.
    vision = SimpleNamespace(extract_image_text=AsyncMock(return_value="Схема электрощита"))
    image = build_image_bytes()

    extracted = await extract_indexable_text("схема.png", image, vision=vision)

    assert extracted == "Схема электрощита"
    assert vision.extract_image_text.await_args.kwargs == {
        "filename": "схема.png",
        "content": image,
    }

    assert await extract_indexable_text(
        "схема.png",
        build_image_bytes(),
        vision=DisabledVisionCapability(),
    ) is None


async def test_truncates_text_above_configured_limit() -> None:
    content = ("строка " * 100).encode("utf-8")

    extracted = await extract_indexable_text(
        "лог.log",
        content,
        vision=DisabledVisionCapability(),
        max_chars=20,
    )

    assert extracted is not None
    assert len(extracted) == 20
